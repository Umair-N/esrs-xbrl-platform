
# File: routes/file_upload_routes.py
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, status
from fastapi.responses import JSONResponse
from typing import List, Optional
import os
import uuid
from datetime import datetime
import PyPDF2
import docx
from io import BytesIO
import mimetypes
from pathlib import Path

# Import from your existing modules
# from core.config import settings
from database import get_db
from model import ReportBlock, ReportDocument, TextUpload
from auth import get_current_user

# Create router
router = APIRouter(prefix="/api/files", tags=["files"])

# Configuration
UPLOAD_DIRECTORY = "uploads"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_DIRECTORY, exist_ok=True)

# Utility functions
def generate_unique_id():
    return str(uuid.uuid4())

def validate_file(file: UploadFile) -> bool:
    """Validate file type and size"""
    if file.size > MAX_FILE_SIZE:
        return False
    
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        return False
    
    return True

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error extracting text from PDF: {str(e)}"
        )

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc_file = BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error extracting text from DOCX: {str(e)}"
        )

def extract_text_from_file(file_content: bytes, file_type: str) -> str:
    """Extract text based on file type"""
    if file_type == "application/pdf":
        return extract_text_from_pdf(file_content)
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                       "application/msword"]:
        return extract_text_from_docx(file_content)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type"
        )

def save_report_to_db(report: ReportDocument, user_id: int, db) -> bool:
    """Save report to database"""
    try:
        from database import create_report
        return create_report(report, user_id, db)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error saving report to database: {str(e)}"
        )

def get_user_reports(user_id: int, db) -> List[ReportDocument]:
    """Get all reports for a user"""
    try:
        from database import get_reports_by_user
        return get_reports_by_user(user_id, db)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error fetching reports: {str(e)}"
        )

def delete_report_from_db(report_id: str, user_id: int, db) -> Optional[str]:
    """Delete a report from database and return file path if exists"""
    try:
        from database import delete_report
        return delete_report(report_id, user_id, db)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error deleting report: {str(e)}"
        )

# API Routes
@router.post("/upload", response_model=ReportDocument)
async def upload_file(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Upload and process a file"""
    
    # Validate file
    if not validate_file(file):
        raise HTTPException(
            status_code=400,
            detail="Invalid file. Please upload a PDF or DOCX file under 10MB."
        )
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Get file type
        file_type = mimetypes.guess_type(file.filename)[0]
        
        # Extract text from file
        extracted_text = extract_text_from_file(file_content, file_type)
        
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the file"
            )
        
        # Save file to disk
        file_id = generate_unique_id()
        file_extension = Path(file.filename).suffix
        saved_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(UPLOAD_DIRECTORY, saved_filename)
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        # Split text into paragraphs
        paragraphs = [p.strip() for p in extracted_text.split('\n\n') if p.strip()]
        
        # Create report document
        report = ReportDocument(
            id=generate_unique_id(),
            title=Path(file.filename).stem,
            created_at=datetime.now().isoformat(),
            updated_at=datetime.now().isoformat(),
            file_path=file_path,
            file_size=len(file_content),
            file_type=file_type,
            blocks=[
                ReportBlock(
                    id=generate_unique_id(),
                    content=paragraph,
                    type="paragraph",
                    tags=[]
                ) for paragraph in paragraphs
            ]
        )
        
        # Save to database
        save_report_to_db(report, current_user["id"], db)
        
        return report
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )

@router.post("/upload-text", response_model=ReportDocument)
async def upload_text(
    text_data: TextUpload,
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Process uploaded text"""
    
    if not text_data.text.strip():
        raise HTTPException(
            status_code=400,
            detail="Text content cannot be empty"
        )
    
    try:
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text_data.text.split('\n\n') if p.strip()]
        
        # Create report document
        report = ReportDocument(
            id=generate_unique_id(),
            title=text_data.title,
            created_at=datetime.now().isoformat(),
            updated_at=datetime.now().isoformat(),
            blocks=[
                ReportBlock(
                    id=generate_unique_id(),
                    content=paragraph,
                    type="paragraph",
                    tags=[]
                ) for paragraph in paragraphs
            ]
        )
        
        # Save to database
        save_report_to_db(report, current_user["id"], db)
        
        return report
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing text: {str(e)}"
        )

@router.get("/reports", response_model=List[ReportDocument])
async def get_reports(
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Get all reports for the current user"""
    return get_user_reports(current_user["id"], db)

@router.get("/reports/{report_id}", response_model=ReportDocument)
async def get_report(
    report_id: str,
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Get a specific report"""
    reports = get_user_reports(current_user["id"], db)
    report = next((r for r in reports if r.id == report_id), None)
    
    if not report:
        raise HTTPException(
            status_code=404,
            detail="Report not found"
        )
    
    return report

@router.delete("/reports/{report_id}")
async def delete_report(
    report_id: str,
    current_user: dict = Depends(get_current_user),
    db = Depends(get_db)
):
    """Delete a report"""
    
    # Delete from database and get file path
    file_path = delete_report_from_db(report_id, current_user["id"], db)
    
    if file_path is None:
        raise HTTPException(
            status_code=404,
            detail="Report not found"
        )
    
    # Delete file if it exists
    if file_path and os.path.exists(file_path):
        os.remove(file_path)
    
    return {"message": "Report deleted successfully"}

# Health check
@router.get("/health")
async def health_check():
    return {"status": "healthy", "service": "file-upload"}